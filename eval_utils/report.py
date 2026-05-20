"""Auto-generated final report for a results directory.

Reads `results/<tag>/<task_id>/{task_result.json, profile.json}` and writes
a `report.md` (always) and `report.docx` (if python-docx is available) into
the same `results/<tag>/` directory.

Sections:
  - Headline counts
  - Per-site breakdown
  - Outcome details (captcha types, stuck reasons)
  - Phase-time distribution
  - Token & cost
  - Slowest / most-stepped
  - Failure-surface (top action errors, DOM drift hot spots)
  - Confounder note (year-in-prompt bucket)
"""
from __future__ import annotations

import json
import re
import statistics as stats
from collections import Counter, defaultdict
from pathlib import Path
from typing import Any, Iterable, Optional


def _load(p: Path) -> dict[str, Any]:
    try:
        return json.loads(p.read_text())
    except Exception:
        return {}


def _percentiles(xs: list[float], qs: tuple[int, ...] = (50, 90, 95)) -> dict[str, float]:
    if not xs:
        return {f"p{q}": 0.0 for q in qs}
    s = sorted(xs)
    return {f"p{q}": s[min(len(s) - 1, max(0, int(round(q / 100 * len(s) - 1))))] for q in qs}


def _money(usd: float) -> str:
    return f"${usd:,.2f}"


def _md_table(headers: list[str], rows: list[list[Any]]) -> str:
    out = ["| " + " | ".join(headers) + " |",
           "|" + "|".join(["---"] * len(headers)) + "|"]
    for r in rows:
        out.append("| " + " | ".join(str(x) for x in r) + " |")
    return "\n".join(out)


def collect(results_dir: Path) -> dict[str, Any]:
    tasks: list[dict[str, Any]] = []
    for task_dir in sorted(p for p in results_dir.iterdir() if p.is_dir()):
        tr = _load(task_dir / "task_result.json")
        if not tr:
            continue
        pr = _load(task_dir / "profile.json")
        tr["_profile"] = pr
        tr["_dir"] = task_dir.name
        tasks.append(tr)
    return {"results_dir": results_dir, "tasks": tasks}


def render(report: dict[str, Any]) -> str:
    tasks: list[dict[str, Any]] = report["tasks"]
    dir_ = Path(report["results_dir"]).name

    n = len(tasks)
    by_outcome = Counter(t.get("success") for t in tasks)
    succ = by_outcome.get("success", 0)
    rate = (succ / n * 100) if n else 0.0

    # ---- aggregate phase times / tokens ----
    wall = [t.get("_profile", {}).get("wall_clock_s") or t.get("duration_seconds") or 0.0 for t in tasks]
    llm_t = [t.get("_profile", {}).get("phase_totals_s", {}).get("llm", 0.0) for t in tasks]
    act_t = [t.get("_profile", {}).get("phase_totals_s", {}).get("browser_action", 0.0) for t in tasks]
    dom_t = [t.get("_profile", {}).get("phase_totals_s", {}).get("dom_extract", 0.0) for t in tasks]
    ss_t = [t.get("_profile", {}).get("phase_totals_s", {}).get("screenshot", 0.0) for t in tasks]
    tok_in = sum((t.get("_profile", {}).get("tokens", {}).get("input") or 0) for t in tasks)
    tok_out = sum((t.get("_profile", {}).get("tokens", {}).get("output") or 0) for t in tasks)

    # ---- per-site ----
    sites: dict[str, list[dict[str, Any]]] = defaultdict(list)
    for t in tasks:
        site = (t.get("task_id") or t.get("_dir") or "").split("--")[0]
        sites[site].append(t)

    site_rows = []
    for site, ts in sorted(sites.items(), key=lambda kv: -len(kv[1])):
        n_s = len(ts)
        s = sum(1 for x in ts if x.get("success") == "success")
        cap = sum(1 for x in ts if x.get("success") == "captcha_blocked")
        st = sum(1 for x in ts if x.get("stuck"))
        steps = [x.get("num_steps", 0) for x in ts]
        dur = [x.get("_profile", {}).get("wall_clock_s") or x.get("duration_seconds", 0.0) for x in ts]
        site_rows.append([
            site, n_s,
            f"{s/n_s*100:.0f}%",
            f"{cap/n_s*100:.0f}%",
            f"{st/n_s*100:.0f}%",
            f"{stats.mean(steps):.1f}" if steps else "-",
            f"{stats.mean(dur):.0f}s" if dur else "-",
        ])

    # ---- outcome detail ----
    stuck_reasons = Counter(t.get("stuck_reason") for t in tasks if t.get("stuck"))
    captcha_types = Counter()
    for t in tasks:
        for ct in (t.get("_profile", {}).get("captcha_types") or []):
            captcha_types[ct] += 1

    # ---- slowest / most-stepped ----
    by_wall = sorted(tasks, key=lambda t: -(t.get("_profile", {}).get("wall_clock_s") or t.get("duration_seconds", 0.0)))[:10]
    by_steps = sorted(tasks, key=lambda t: -(t.get("num_steps") or 0))[:10]

    # ---- error surface ----
    err_counter: Counter[str] = Counter()
    for t in tasks:
        # parse steps.jsonl if present for first-error of each step
        sf = (Path(report["results_dir"]) / t["_dir"] / "steps.jsonl")
        if sf.exists():
            for line in sf.read_text().splitlines():
                try:
                    e = json.loads(line).get("error")
                except Exception:
                    e = None
                if e:
                    err_counter[re.sub(r"\s+", " ", e)[:120]] += 1

    # ---- year-in-prompt bucket (confounder note) ----
    yr_pat = re.compile(r"\b(20[12]\d)\b")
    by_year: dict[str, Counter[str]] = defaultdict(Counter)
    for t in tasks:
        yrs = yr_pat.findall(t.get("task_prompt") or "")
        bucket = "no year"
        if yrs:
            y_max = max(int(y) for y in yrs)
            if y_max >= 2027: bucket = "≥2027"
            elif y_max >= 2026: bucket = "2026"
            else: bucket = "≤2025"
        by_year[bucket][t.get("success") or "?"] += 1

    # ---- render md ----
    md = []
    md.append(f"# browser-use-eval report — `{dir_}`\n")
    md.append(f"- total tasks: **{n}**")
    md.append(f"- success: **{succ}**  (rate {rate:.1f}%)")
    for k in ("failed", "captcha_blocked", "stuck", "unknown"):
        md.append(f"- {k}: **{by_outcome.get(k, 0)}**")
    md.append(f"- wall clock total: **{sum(wall)/3600:.1f}h**")
    md.append(f"- tokens: **in {tok_in:,}** / **out {tok_out:,}**\n")

    md.append("## Per-site\n")
    md.append(_md_table(
        ["site", "N", "success%", "captcha%", "stuck%", "avg steps", "avg duration"], site_rows))

    if stuck_reasons:
        md.append("\n## Stuck reasons\n")
        md.append(_md_table(["reason", "count"], list(stuck_reasons.items())))

    if captcha_types:
        md.append("\n## CAPTCHA types encountered\n")
        md.append(_md_table(["type", "count"], list(captcha_types.items())))

    md.append("\n## Phase timing (seconds)\n")
    ph_rows = []
    for label, xs in [("LLM inference", llm_t), ("Browser action", act_t),
                      ("DOM extract", dom_t), ("Screenshot", ss_t),
                      ("Wall clock", wall)]:
        p = _percentiles(xs)
        ph_rows.append([label, f"{stats.mean(xs):.1f}" if xs else "-",
                        f"{p['p50']:.1f}", f"{p['p90']:.1f}", f"{p['p95']:.1f}"])
    md.append(_md_table(["phase", "mean", "p50", "p90", "p95"], ph_rows))

    md.append("\n## Top 10 slowest tasks\n")
    md.append(_md_table(
        ["task_id", "outcome", "steps", "wall_s"],
        [[t.get("task_id"), t.get("success"), t.get("num_steps"),
          f"{(t.get('_profile', {}).get('wall_clock_s') or t.get('duration_seconds', 0.0)):.0f}"]
         for t in by_wall]))

    md.append("\n## Top 10 most-stepped tasks\n")
    md.append(_md_table(
        ["task_id", "outcome", "steps", "wall_s"],
        [[t.get("task_id"), t.get("success"), t.get("num_steps"),
          f"{(t.get('_profile', {}).get('wall_clock_s') or t.get('duration_seconds', 0.0)):.0f}"]
         for t in by_steps]))

    if err_counter:
        md.append("\n## Top 20 action errors\n")
        md.append(_md_table(["count", "error"],
                            [[c, msg] for msg, c in err_counter.most_common(20)]))

    md.append("\n## Confounder: outcome by year mentioned in prompt\n")
    md.append(_md_table(
        ["year bucket", "success", "failed", "captcha_blocked", "stuck", "unknown"],
        [[b, by_year[b].get("success", 0), by_year[b].get("failed", 0),
          by_year[b].get("captcha_blocked", 0), by_year[b].get("stuck", 0),
          by_year[b].get("unknown", 0)] for b in ("no year", "≤2025", "2026", "≥2027")]))

    return "\n".join(md) + "\n"


def write(results_dir: Path) -> Path:
    """Build and write report.md (+report.docx if python-docx available)."""
    report = collect(results_dir)
    md = render(report)
    out_md = results_dir / "report.md"
    out_md.write_text(md)
    try:
        from docx import Document  # type: ignore
        doc = Document()
        for line in md.splitlines():
            if line.startswith("# "):    doc.add_heading(line[2:], level=1)
            elif line.startswith("## "): doc.add_heading(line[3:], level=2)
            else:                        doc.add_paragraph(line)
        doc.save(results_dir / "report.docx")
    except Exception:
        pass
    return out_md


if __name__ == "__main__":
    import sys
    p = Path(sys.argv[1]) if len(sys.argv) > 1 else Path("results/examples-browser-use-vision-true")
    print("wrote", write(p))
